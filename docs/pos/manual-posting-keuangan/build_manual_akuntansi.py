from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path(r"C:\opt\AIS\ais\src\main\docs\pos\manual-posting-keuangan")
SHOT = BASE / "screenshots-uat"
OUTPUT = BASE / "Manual-UAT-Akuntansi-eBisnis.docx"

NAVY = "132238"
BLUE = "246B97"
SKY = "DCEEF8"
LIGHT = "F3F6FA"
MID = "D9E2EC"
TEXT = "172033"
MUTED = "60708A"
GREEN = "E6F4EA"
AMBER = "FFF3D6"
RED = "FCE8E6"
WHITE = "FFFFFF"


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, *, bold=False, color=TEXT, size=8.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_with_next(paragraph, keep=True):
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:keepNext"))
    if keep and node is None:
        p_pr.append(OxmlElement("w:keepNext"))
    elif not keep and node is not None:
        p_pr.remove(node)


def page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Halaman ")
    run.font.size = Pt(8)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def configure_document(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(1.55)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)
    sec.header_distance = Cm(0.65)
    sec.footer_distance = Cm(0.65)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in (
        ("Title", 30, NAVY),
        ("Heading 1", 19, NAVY),
        ("Heading 2", 13.5, BLUE),
        ("Heading 3", 10.5, NAVY),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.keep_with_next = True

    header = sec.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.7))
    table.autofit = False
    table.columns[0].width = Cm(8.85)
    table.columns[1].width = Cm(8.85)
    set_cell_text(table.cell(0, 0), "eBisnis POS  •  AKUNTANSI", bold=True, color=BLUE, size=8)
    set_cell_text(table.cell(0, 1), "Manual UAT & Panduan Pengguna", bold=True, color=MUTED, size=8,
                  align=WD_ALIGN_PARAGRAPH.RIGHT)
    for c in table.rows[0].cells:
        margins(c, 0, 0, 40, 0)

    footer = sec.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Dokumen UAT • Data demo • 4 September 2026")
    r.font.name = "Aptos"
    r.font.size = Pt(8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(16), WD_ALIGN_PARAGRAPH.RIGHT)
    p.add_run("\t")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_band(doc, text: str, fill=SKY, color=NAVY, size=10.5):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_text(cell, text, bold=True, color=color, size=size)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc, label: str, text: str, tone="info"):
    fill = {"info": SKY, "ok": GREEN, "warn": AMBER, "risk": RED}.get(tone, SKY)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(14.8)
    shade(table.cell(0, 0), fill)
    shade(table.cell(0, 1), fill)
    set_cell_text(table.cell(0, 0), label.upper(), bold=True, color=NAVY, size=8)
    set_cell_text(table.cell(0, 1), text, color=TEXT, size=8.7)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_steps(doc, steps):
    for idx, step in enumerate(steps, 1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Cm(0.15)
        p.paragraph_format.first_line_indent = Cm(-0.15)
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(f"{idx}. ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(BLUE)
        p.add_run(step)


def add_screenshot(doc, filename: str, caption: str, *, width=6.7):
    path = SHOT / filename
    if not path.exists():
        add_callout(doc, "Bukti tertunda", f"Tangkapan layar {filename} belum tersedia saat dokumen dibangun.", "warn")
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.paragraph_format.space_after = Pt(5)
    r = c.add_run(caption)
    r.italic = True
    r.font.name = "Aptos"
    r.font.size = Pt(7.8)
    r.font.color.rgb = RGBColor.from_string(MUTED)
    return True


def add_page(doc, title, intro, shot, caption, steps=None, callout=None):
    doc.add_heading(title, level=1)
    if intro:
        doc.add_paragraph(intro)
    if steps:
        add_steps(doc, steps)
    add_screenshot(doc, shot, caption)
    if callout:
        add_callout(doc, *callout)
    doc.add_page_break()


def add_table(doc, headers, rows, widths=None, font_size=7.7):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False if widths else True
    if widths:
        # Word can silently widen columns when only table.columns is set.  A
        # fixed layout plus explicit cell widths keeps wide tables inside A4.
        max_width = 16.5
        scale = min(1.0, max_width / sum(widths))
        widths = [width * scale for width in widths]
        tbl_pr = table._tbl.tblPr
        tbl_layout = tbl_pr.find(qn("w:tblLayout"))
        if tbl_layout is None:
            tbl_layout = OxmlElement("w:tblLayout")
            tbl_pr.append(tbl_layout)
        tbl_layout.set(qn("w:type"), "fixed")
        for i, width in enumerate(widths):
            table.columns[i].width = Cm(width)

    def set_widths(cells):
        if not widths:
            return
        for i, width in enumerate(widths):
            cells[i].width = Cm(width)
            tc_pr = cells[i]._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(round(width / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")

    set_widths(table.rows[0].cells)
    for i, h in enumerate(headers):
        shade(table.cell(0, i), NAVY)
        set_cell_text(table.cell(0, i), h, bold=True, color=WHITE, size=8)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        set_widths(cells)
        for i, value in enumerate(row):
            if len(table.rows) % 2 == 0:
                shade(cells[i], LIGHT)
            set_cell_text(cells[i], value, size=font_size)
    return table


def cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    tag = doc.add_table(rows=1, cols=1)
    shade(tag.cell(0, 0), SKY)
    set_cell_text(tag.cell(0, 0), "BUILD VARIAN eBISNIS  •  UAT 04-09-2026", bold=True, color=BLUE, size=9)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(34)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("MANUAL UAT")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor.from_string(BLUE)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(10)
    p.add_run("Akuntansi eBisnis")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Dari Jurnal Umum, proses posting, siklus akuntansi, master data, hingga laporan keuangan")
    r.font.name = "Aptos Display"
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor.from_string(MUTED)

    add_screenshot(doc, "01-menu-akuntansi-terbuka.png",
                   "Menu Akuntansi pada aplikasi eBisnis POS versi 1.34.20 (build 182).",
                   width=6.55)

    table = doc.add_table(rows=4, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(12.8)
    for i, (k, v) in enumerate((
        ("Produk", "eBisnis POS — varian eBisnis"),
        ("Versi diuji", "1.34.20 (build 182)"),
        ("Lingkungan", "Server demo eBisnis • Kantin Demo • Semua Toko"),
        ("Keluaran", "Panduan pengguna + hasil UAT berbukti screenshot"),
    )):
        shade(table.cell(i, 0), NAVY)
        set_cell_text(table.cell(i, 0), k, bold=True, color=WHITE, size=8)
        shade(table.cell(i, 1), LIGHT)
        set_cell_text(table.cell(i, 1), v, size=8.4)
    doc.add_page_break()


def build():
    doc = Document()
    configure_document(doc)
    cover(doc)

    doc.add_heading("Tentang dokumen ini", level=1)
    doc.add_paragraph(
        "Dokumen ini menggabungkan panduan operasional dan hasil User Acceptance Test (UAT) "
        "menu Akuntansi. Seluruh gambar berasal dari render aplikasi eBisnis yang benar-benar "
        "terhubung ke server demo; bukan mockup. Data contoh mencakup jurnal manual, saldo awal, "
        "penyesuaian, penjualan/kulakan, pembayaran hutang, penerimaan piutang, dan anggaran."
    )
    add_callout(doc, "Keamanan", "Kata sandi demo tidak dicantumkan. Gunakan kredensial yang disampaikan melalui kanal terpisah.", "info")
    add_callout(doc, "Batas UAT", "Tutup Buku dan Closing tidak dijalankan karena keduanya mengunci periode bersama. Layar dan pratinjaunya tetap diuji.", "warn")
    doc.add_heading("Ruang lingkup", level=2)
    add_table(doc, ["Area", "Yang diuji", "Bukti"], [
        ("Akses", "Login, identitas admin, Kantin Demo, menu Akuntansi", "Gambar 1–2"),
        ("Jurnal Umum", "Buat draf seimbang, simpan, konfirmasi, posting massal", "Gambar 3–8"),
        ("Posting otomatis", "HPP, Penjualan, Kulakan, Bayar Hutang, Terima Piutang", "Gambar 10–14"),
        ("Siklus", "Saldo Awal, Penyesuaian, Tutup Buku, Closing", "Gambar 15–18"),
        ("Master", "Anggaran, Kode Akun, Grup, Jenis Transaksi, Bank", "Gambar 20–24"),
        ("Laporan", "Katalog dan enam laporan inti; delapan endpoint diverifikasi", "Gambar 19, 25–30"),
    ], widths=[3.2, 9.4, 4.0])
    doc.add_heading("Konvensi status", level=2)
    add_table(doc, ["Status", "Makna"], [
        ("Lulus", "Layar memuat, aksi selesai, dan hasil sesuai."),
        ("Lulus dengan catatan", "Fungsi utama tampil, tetapi data/prasyarat atau tata letak perlu tindak lanjut."),
        ("Tidak dieksekusi", "Aksi berisiko mengunci periode bersama; panduan disediakan tanpa menekan aksi final."),
    ], widths=[4.0, 12.6])
    doc.add_page_break()

    doc.add_heading("Ringkasan hasil UAT", level=1)
    rows = [
        ("Draft Jurnal", "Lulus", "52 aktivitas: 39 draf, 13 terposting, 0 closing."),
        ("Jurnal Umum", "Catatan", "8 jurnal manual terposting; daftar sempat lulus, tetapi endpoint daftar regresi setelah restart."),
        ("Posting HPP", "Catatan", "1 transaksi Rp1.200.000; akun HPP ada, tetapi akun persediaan Master Aset masih kosong."),
        ("Posting Penjualan", "Catatan", "Perpindahan submenu dapat mempertahankan pratinjau HPP; indikasi state layar tidak direset."),
        ("Posting Kulakan", "Catatan", "2 faktur Rp600.000 dan Rp450.000 tampil; tertahan oleh akun persediaan Master Aset."),
        ("Posting Bayar Hutang", "Lulus", "1 pembayaran Rp150.000 siap; 1 pembayaran Rp200.000 telah diposting saat penyiapan."),
        ("Posting Terima Piutang", "Lulus", "1 penerimaan Rp200.000 siap; 1 penerimaan Rp175.000 telah diposting saat penyiapan."),
        ("Saldo Awal", "Lulus", "5 akun terposting; total Debet dan Kredit masing-masing Rp22.000.000."),
        ("Jurnal Penyesuaian", "Lulus", "3 template sampel: listrik, sewa outlet, dan cadangan kerugian piutang."),
        ("Tutup Buku", "Tidak dieksekusi", "Draf tampil; layar tetap menyatakan Laba Ditahan belum diatur meski toko sudah dipetakan."),
        ("Closing", "Tidak dieksekusi", "Form Closing Baru tampil; penyimpanan final sengaja tidak dijalankan."),
        ("Katalog Laporan", "Lulus", "9 kategori laporan keuangan/pengadaan/pajak/anggaran."),
        ("Anggaran", "Catatan", "3 item, total rencana Rp45.000.000; ada overflow horizontal pada lebar 1264 px."),
        ("Kode Akun", "Lulus", "317 akun tampil; 256 akun yang belum terpetakan telah ditautkan secara aman."),
        ("Grup Akun", "Lulus", "Grup dan jumlah akun tampil."),
        ("Jenis Transaksi", "Lulus", "Kode jurnal AJE/GJP/JBI/JBO/JCI/JCO/JMM/PB tampil."),
        ("Bank", "Lulus", "Master bank sample tampil."),
    ]
    add_table(doc, ["Submenu", "Status", "Hasil ringkas"], rows, widths=[4.0, 3.0, 9.6], font_size=7.2)
    doc.add_paragraph()
    add_callout(doc, "Kesimpulan", "Data sampel kini mengisi seluruh tahap utama dan enam laporan inti. Posting bayar hutang/terima piutang siap; HPP, Penjualan, Kulakan, Tutup Buku, dan endpoint daftar Jurnal Umum masih memiliki temuan yang perlu diperbaiki.", "warn")
    doc.add_page_break()

    add_page(doc, "1. Masuk dan memastikan konteks", "Gunakan akun demo yang diberikan terpisah. Setelah masuk, pastikan identitas admin, toko, dan versi build sesuai.",
             "00-layar-awal-integration.png", "Gambar 1. Beranda POS setelah login pada Kantin Demo.",
             ["Buka build varian eBisnis.", "Masukkan akun demo, lalu pilih Masuk.", "Pastikan nama pengguna admin dan Kantin Demo tampil.", "Pastikan versi 1.34.20 (build 182) terlihat pada sidebar."],
             ("Kontrol", "Jangan melanjutkan jurnal bila toko atau periode yang tampil tidak sesuai.", "warn"))

    add_page(doc, "2. Membuka seluruh menu Akuntansi", "Klik grup AKUNTANSI pada sidebar. Terdapat 17 submenu yang diuji.",
             "01-menu-akuntansi-terbuka.png", "Gambar 2. Seluruh submenu Akuntansi.",
             ["Gulir sidebar sampai menemukan AKUNTANSI.", "Klik judul grup agar daftar terbuka.", "Pilih submenu berdasarkan proses yang hendak dilakukan."],
             ("Catatan navigasi", "Pada sesi UAT, submenu yang berbagi layar bertab kadang mempertahankan tab sebelumnya. Bila header sudah berubah tetapi isi belum, klik tab target pada bagian atas layar.", "warn"))

    add_page(doc, "3. Daftar Jurnal Umum", "Daftar menampilkan jurnal contoh beserta kode, tanggal, keterangan, total, status, dan aksi. Bukti ini diambil sebelum restart server; pengujian ulang setelah restart menemukan regresi internal pada endpoint daftar.",
             "02-jurnal-umum-daftar.png", "Gambar 3. Daftar Jurnal Umum pada periode uji.",
             ["Pilih Akuntansi → Jurnal Umum.", "Atur Mulai dan Sampai.", "Gunakan pencarian atau filter Status bila diperlukan.", "Klik Jurnal Baru untuk memasukkan transaksi manual."])

    add_page(doc, "4. Membuat jurnal baru", "Contoh pertama adalah setoran modal awal: Kas Yayasan didebet dan Modal dikredit.",
             "03-jurnal-umum-formulir-baru.png", "Gambar 4. Form Jurnal Umum Baru.",
             ["Pilih tanggal transaksi.", "Isi keterangan yang mudah ditelusuri.", "Pilih akun pada setiap baris.", "Isi hanya salah satu sisi—Debet atau Kredit—pada satu baris."])

    add_page(doc, "5. Memastikan jurnal seimbang", "Jurnal hanya dapat disimpan bila minimal dua baris sah dan jumlah Debet sama dengan Kredit.",
             "04-jurnal-umum-seimbang.png", "Gambar 5. Jurnal sample seimbang dan berstatus Siap disimpan.",
             ["Debet 111.101 Kas Yayasan sebesar Rp1.000.000.", "Kredit 800.000 Modal sebesar Rp1.000.000.", "Pastikan chip Siap disimpan muncul.", "Klik Simpan sebagai Draf."],
             ("Praktik baik", "Gunakan uraian kepala dan uraian baris yang menjelaskan sumber transaksi, dokumen pendukung, serta tujuan dana.", "info"))

    add_page(doc, "6. Memeriksa jurnal tersimpan", "UAT menggunakan dua contoh: setoran modal awal dan pemindahan dana kas kecil. Setelah proses posting, keduanya ditampilkan berstatus Terposting.",
             "05-jurnal-umum-draf-tersimpan.png", "Gambar 6. Dua jurnal contoh pada periode uji setelah posting.",
             ["Cari jurnal menggunakan kata UAT Manual Akuntansi.", "Cocokkan Debet dan Kredit pada setiap jurnal.", "Sebelum posting, pastikan statusnya Draf; sesudah posting, pastikan statusnya Terposting.", "Buka menu aksi jika perlu menelusuri rinciannya."],
             ("Kontrol empat mata", "Untuk operasional nyata, pemilik dokumen dan petugas posting sebaiknya berbeda.", "info"))

    add_page(doc, "7. Konfirmasi posting", "Posting memindahkan jurnal ke buku besar dan menjadikannya sumber laporan keuangan.",
             "06-jurnal-umum-konfirmasi-posting.png", "Gambar 7. Dialog Posting Semua Draf.",
             ["Klik Posting Semua Draf pada periode terpilih.", "Baca jumlah jurnal yang akan diproses.", "Klik Posting untuk melanjutkan atau Batal untuk memeriksa kembali."],
             ("Dampak", "Jurnal terposting tidak dapat diedit langsung. Batalkan posting lebih dahulu bila masih diizinkan dan periode belum closing.", "warn"))

    add_page(doc, "8. Memverifikasi hasil posting", "Status kedua jurnal berubah menjadi Terposting. Nilai Debet dan Kredit tetap seimbang.",
             "07-jurnal-umum-terposting.png", "Gambar 8. JU/09/00001 dan JU/09/00002 telah terposting.",
             ["Pastikan status Terposting muncul.", "Cocokkan kode jurnal dan nilai.", "Lanjutkan ke Katalog Laporan untuk menelusuri dampaknya pada buku besar dan neraca."],
             ("Hasil UAT", "JU/09/00001 Rp1.000.000 dan JU/09/00002 Rp250.000 berhasil diposting pada 4 September 2026.", "ok"))

    add_page(doc, "9. Draft Jurnal lintas modul", "Layar ini merangkum kesiapan jurnal dari seluruh modul sebelum closing.",
             "08-draft-jurnal.png", "Gambar 9. Ringkasan Draft Jurnal: 39 draf, 13 terposting, 0 closing, total 52 aktivitas.",
             ["Pilih Akuntansi → Draft Jurnal.", "Atur rentang tanggal.", "Periksa kartu Draft, Terposting, Closing, dan Total Aktivitas.", "Buka detail per jenis jurnal melalui menu aksi."],
             ("UAT", "Delapan jurnal manual terposting, saldo awal, penyesuaian, dan aktivitas lintas modul ikut membentuk ringkasan.", "ok"))

    posting_pages = [
        ("10. Posting HPP", "09-posting-hpp.png", "Gambar 10. Pratinjau Posting HPP.",
         "Pratinjau menemukan satu transaksi ABC Kecap Manis Rp1.200.000. Akun HPP 510.900 sudah terbaca, tetapi akun lawan Persediaan pada Master Aset belum tersedia."),
        ("11. Posting Penjualan", "10-posting-penjualan.png", "Gambar 11. Pratinjau Posting Penjualan.",
         "Saat berpindah dari Posting HPP, layar mempertahankan pratinjau HPP. Ini adalah defect state layar; jangan menekan Posting sebelum judul dan isi sama-sama menunjukkan Penjualan."),
        ("12. Posting Kulakan", "11-posting-kulakan.png", "Gambar 12. Posting Kulakan.",
         "Dua faktur pemasok tampil: Rp600.000 dan Rp450.000. Keduanya belum siap karena akun persediaan Master Aset belum tersedia."),
        ("13. Posting Bayar Hutang", "12-posting-bayar-hutang.png", "Gambar 13. Posting Bayar Hutang.",
         "Satu pembayaran Rp150.000 siap diposting (Debet 310.600; Kredit 111.101); satu pembayaran Rp200.000 telah diposting saat penyiapan data."),
        ("14. Posting Terima Piutang", "13-posting-terima-piutang.png", "Gambar 14. Posting Terima Piutang.",
         "Satu penerimaan Rp200.000 siap diposting (Debet 111.101; Kredit 131.300); satu penerimaan Rp175.000 telah diposting saat penyiapan data."),
    ]
    for title, shot, cap, outcome in posting_pages:
        add_page(doc, title, outcome, shot, cap,
                 ["Pilih periode transaksi.", "Klik Muat ulang/Pratinjau.", "Periksa baris yang siap dan alasan baris tertahan.", "Lengkapi pemetaan akun bila diperlukan.", "Posting hanya baris yang siap."],
                 ("Status UAT", outcome, "warn" if "belum" in outcome or "defect" in outcome else "info"))

    cycle_pages = [
        ("15. Saldo Awal (Neraca Awal)", "14-saldo-awal.png", "Gambar 15. Lima saldo awal telah diposting.",
         "Lima akun pembukaan telah diposting: kas, persediaan, piutang, utang, dan modal. Total Debet dan Kredit masing-masing Rp22.000.000."),
        ("16. Jurnal Penyesuaian Berkala", "15-jurnal-penyesuaian.png", "Gambar 16. Template Jurnal Penyesuaian.",
         "Tiga template sampel tersedia untuk listrik/internet Rp350.000, sewa outlet Rp500.000, dan cadangan kerugian piutang Rp100.000; jurnal periode September telah diposting."),
        ("17. Tutup Buku (Laba Ditahan)", "16-tutup-buku.png", "Gambar 17. Pratinjau Tutup Buku.",
         "Gunakan Lihat Draf untuk meninjau akun laba/rugi dan sisi penutupnya. UAT menemukan pesan Laba Ditahan belum diatur walau toko sudah mempunyai pemetaan; aksi final tidak dijalankan."),
        ("18. Closing", "17-closing.png", "Gambar 18. Form Closing Baru.",
         "Form Closing Baru berhasil dibuka. Closing menautkan jurnal sampai tanggal batas dan mencegah pembatalan posting; formulir ditutup tanpa menyimpan pada database demo bersama."),
    ]
    for title, shot, cap, intro in cycle_pages:
        tone = "warn" if "Tutup" in title or "Closing" in title else "info"
        add_page(doc, title, intro, shot, cap,
                 ["Buka submenu yang diperlukan.", "Klik tab target bila isi masih berada pada tab sebelumnya.", "Lengkapi parameter/periode.", "Tinjau draf sebelum aksi final."],
                 ("Kontrol periode", "Cadangkan data dan pastikan seluruh jurnal telah direkonsiliasi sebelum Tutup Buku atau Closing.", tone))

    add_page(doc, "19. Katalog Laporan Keuangan", "Katalog menampilkan laporan berbasis jurnal sebagai laporan resmi, serta laporan operasional sebagai pembanding.",
             "18-katalog-laporan.png", "Gambar 19. Katalog Laporan Keuangan.",
             ["Pilih Akuntansi → Katalog Laporan.", "Gunakan kategori atau pencarian.", "Pilih laporan berlabel Data.", "Atur periode dan klik Tampilkan."],
             ("Prioritas", "Untuk laporan resmi gunakan judul yang memuat “Berbasis Jurnal Akuntansi”; angka hanya mengambil jurnal yang sudah terposting.", "info"))

    add_page(doc, "20. Anggaran (RAB Bulanan)", "Pilih tahun, satuan kerja, sumber dana, dan revisi untuk melihat rencana, realisasi, serta penggunaan anggaran.",
             "19-anggaran.png", "Gambar 20. Tiga item RAB Bulanan dengan total rencana Rp45.000.000.",
             ["Pilih filter anggaran.", "Klik Terapkan.", "Gunakan tab Rencana Bulanan, Realisasi, atau Penggunaan Anggaran.", "Tambah Item atau Buat Revisi Baru sesuai hak akses."],
             ("Temuan tata letak", "Pada lebar render 1264 px, kalender bulan melampaui area tabel (RenderFlex overflow 370 px). Data tetap dapat dibaca sebagian; perbaikan responsif diperlukan.", "warn"))

    master_pages = [
        ("21. Kode Akun", "20-kode-akun.png", "Gambar 21. Bagan akun sample—317 akun.",
         "Kelola struktur akun, saldo normal, klasifikasi laporan, dan pemakaian akun."),
        ("22. Grup Akun", "21-grup-akun.png", "Gambar 22. Grup akun dan jumlah anggotanya.",
         "Gunakan grup untuk klasifikasi dan penyajian laporan."),
        ("23. Jenis Transaksi", "22-jenis-transaksi.png", "Gambar 23. Jenis transaksi dan kode jurnal.",
         "Atur kode penomoran serta akun bawaan untuk transaksi berulang."),
        ("24. Bank", "23-bank.png", "Gambar 24. Master Bank.",
         "Hubungkan bank atau kas dengan akun kas/bank yang digunakan jurnal."),
    ]
    for title, shot, cap, intro in master_pages:
        add_page(doc, title, intro, shot, cap,
                 ["Pilih submenu.", "Jika isi masih menampilkan tab lain, klik tab target.", "Gunakan pencarian dan Terapkan.", "Tambah/ubah data hanya setelah memeriksa kode serta relasi akunnya."],
                 ("Best practice", "Jangan menghapus akun yang sudah dipakai transaksi. Nonaktifkan atau gunakan mekanisme Data Terhapus bila tersedia.", "info"))

    doc.add_heading("25. Laporan keuangan", level=1)
    doc.add_paragraph(
        "Enam laporan inti diuji ulang untuk periode 1–30 September 2026 setelah 256 akun lama "
        "ditautkan ke Kelompok Laporan melalui pemetaan standar aplikasi. Semua layar memuat dan "
        "tombol PDF/Excel aktif."
    )
    add_table(doc, ["Laporan", "Bukti data sampel", "Hasil"], [
        ("Laba Rugi", "Pendapatan Rp4.750.000; beban berasal dari HPP, listrik, sewa, dan penyisihan.", "Sukses"),
        ("Neraca", "Kas Yayasan Rp18.025.000 serta akun aset/kewajiban/modal lain tampil terkelompok.", "Sukses"),
        ("Arus Kas", "Saldo awal kas/bank Rp12.900.000; penerimaan terlihat Rp3.100.000.", "Sukses"),
        ("Keseluruhan Jurnal", "JU/09/00001 s.d. JU/09/00008 dapat ditelusuri per baris.", "Sukses"),
        ("Buku Besar", "Mutasi Kas Yayasan dan akun lawan tampil per nomor jurnal.", "Sukses"),
        ("Neraca Saldo", "Kas, piutang, persediaan, penyisihan, sewa, dan akun lain memiliki saldo.", "Sukses"),
    ], widths=[4.3, 9.8, 2.5])
    add_callout(doc, "Rekonsiliasi", "Pemetaan Kelompok Laporan tidak mengubah saldo atau jurnal; ia menentukan tempat akun disajikan pada laporan resmi.", "info")
    doc.add_page_break()

    report_pages = [
        ("25.1 Laba Rugi", "24-laporan-laba-rugi.png", "Gambar 25. Laba Rugi berbasis jurnal.",
         "Pendapatan penjualan sampel Rp4.750.000 tampil setelah akun 410.900 dipetakan. Beban sampel mencakup HPP, listrik/internet, sewa outlet, dan cadangan kerugian piutang."),
        ("25.2 Neraca", "25-laporan-neraca.png", "Gambar 26. Neraca berbasis jurnal.",
         "Kas Yayasan Rp18.025.000 serta akun neraca lain tampil pada kelompok yang sesuai; label 'belum dipetakan' tidak lagi mendominasi hasil."),
        ("25.3 Arus Kas", "26-laporan-arus-kas.png", "Gambar 27. Arus Kas berbasis jurnal.",
         "Saldo awal Kas & Bank Rp12.900.000 dan penerimaan Rp3.100.000 terlihat; akun lawan memudahkan penelusuran sumber arus."),
        ("25.4 Keseluruhan Jurnal", "27-laporan-jurnal-umum.png", "Gambar 28. Keseluruhan Jurnal.",
         "Delapan jurnal manual terposting dapat ditelusuri, termasuk penjualan tunai/kredit, HPP, listrik, sewa, dan penyisihan."),
        ("25.5 Buku Besar", "28-laporan-buku-besar.png", "Gambar 29. Rincian Buku Besar.",
         "Mutasi Kas Yayasan dapat ditelusuri ke setoran modal, pemindahan kas, penjualan tunai, dan pembayaran biaya."),
        ("25.6 Neraca Saldo", "29-laporan-neraca-saldo.png", "Gambar 30. Neraca Percobaan / Neraca Saldo.",
         "Saldo akun sampel tampil, antara lain Kas Rp3.375.000, Piutang Usaha Rp1.075.000, Cadangan Kerugian Piutang Rp200.000 (kredit), dan Persediaan Rp2.100.000 (kredit) untuk periode yang dipilih."),
    ]
    for title, shot, cap, hasil_uat in report_pages:
        add_page(doc, title, "Atur periode 1–30 September 2026 dan pilih Semua Unit/Konsolidasi bila diperlukan, lalu klik Tampilkan.",
                 shot, cap,
                 ["Periksa judul dan periode.", "Pastikan tabel hasil terisi.", "Gunakan PDF atau Excel setelah hasil tampil.", "Telusuri angka melalui fitur rincian/asal angka bila tersedia."],
                 ("UAT", hasil_uat, "ok"))

    doc.add_heading("26. Temuan dan rekomendasi", level=1)
    add_table(doc, ["ID", "Prioritas", "Temuan", "Rekomendasi"], [
        ("ACC-01", "Tinggi", "Posting HPP dan Kulakan tertahan karena Master Aset produk belum menyediakan akun Persediaan.", "Lengkapi akun Persediaan pada Master/Kelompok Aset yang benar, lalu ulang pratinjau sebelum posting."),
        ("ACC-02", "Tinggi", "Posting Penjualan dapat mempertahankan isi pratinjau HPP setelah perpindahan submenu.", "Reset state pada didUpdateWidget atau gunakan key unik; blokir posting bila jenis layar dan payload berbeda."),
        ("ACC-03", "Sedang", "Tutup Buku menyatakan akun Laba Ditahan belum diatur walau akun sudah tersimpan pada toko.", "Pastikan tokoId/konteks toko ikut dikirim saat pratinjau dan gunakan pemetaan toko aktif."),
        ("ACC-04", "Sedang", "RAB Bulanan overflow 370 px pada lebar render 1264 px.", "Buat header bulan horizontal-scroll atau responsif tanpa RenderFlex overflow."),
        ("API-01", "Tinggi", "Setelah restart, jurnal_umum_daftar mengembalikan SERVER_ERROR walau laporan jurnal tetap dapat memuat data.", "Telusuri referensi API-MTLWU4YI dan uji endpoint daftar setelah perbaikan server."),
        ("ENV-01", "Info", "Server sempat direstart; data dan laporan kemudian diuji ulang setelah layanan aktif.", "Jalankan health check API sebelum regression test dan catat window maintenance."),
    ], widths=[2.0, 2.2, 6.2, 6.2], font_size=7.2)
    doc.add_heading("Keputusan kesiapan", level=2)
    add_callout(doc, "Siap bersyarat", "Data master, saldo awal, penyesuaian, posting hutang/piutang, dan laporan dapat digunakan. Tahan Posting HPP/Penjualan/Kulakan dan Tutup Buku sampai ACC-01–ACC-03 serta API-01 diperbaiki; Closing final perlu UAT pada salinan database.", "warn")
    doc.add_page_break()

    doc.add_heading("27. Checklist operasional", level=1)
    add_table(doc, ["Tahap", "Pemeriksaan", "Selesai"], [
        ("Sebelum input", "Toko, unit, periode, dan mata uang benar.", "☐"),
        ("Draf", "Minimal dua baris; tiap baris hanya satu sisi.", "☐"),
        ("Validasi", "Total Debet = total Kredit; dokumen pendukung tersedia.", "☐"),
        ("Posting", "Status berubah menjadi Terposting; kode jurnal tercatat.", "☐"),
        ("Rekonsiliasi", "Buku Besar, Neraca Saldo, dan laporan sumber cocok.", "☐"),
        ("Periode", "Tidak ada draf penting sebelum Tutup Buku/Closing.", "☐"),
        ("Ekspor", "PDF/Excel diberi nama periode dan disimpan di lokasi resmi.", "☐"),
    ], widths=[3.2, 11.8, 1.6])
    doc.add_heading("Alur singkat", level=2)
    add_band(doc, "BUAT DRAF  →  VALIDASI DEBET=KREDIT  →  POSTING  →  BUKU BESAR  →  NERACA SALDO  →  LAPORAN", SKY, NAVY, 10)
    doc.add_paragraph(
        "Jika angka laporan belum sesuai, kembali ke daftar jurnal dan pastikan transaksi sudah terposting, "
        "periode/unit benar, serta akun telah masuk klasifikasi laporan yang tepat."
    )
    add_callout(doc, "Akhir dokumen", "Manual ini dibuat dari UAT server demo pada 4 September 2026 dan pengujian ulang setelah restart. Ulangi smoke test setelah perbaikan ACC-01 sampai ACC-04 dan API-01.", "info")

    props = doc.core_properties
    props.title = "Manual UAT Akuntansi eBisnis"
    props.subject = "Jurnal Umum, posting, siklus akuntansi, master data, dan laporan keuangan"
    props.author = "Tim UAT eBisnis"
    props.keywords = "eBisnis, Akuntansi, UAT, Jurnal Umum, Posting, Laporan Keuangan"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
