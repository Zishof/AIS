from __future__ import annotations

import importlib.util
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor


BASE = Path(r"C:\opt\AIS\ais\src\main\docs\pos\manual-keuangan-akuntansi")
SHOT = BASE / "screenshots-uat"
OUTPUT = BASE / "Manual-UAT-Keuangan-dan-Integrasi-Akuntansi-eBisnis.docx"
HELPER = BASE.parent / "manual-posting-keuangan" / "build_manual_akuntansi.py"

spec = importlib.util.spec_from_file_location("manual_akuntansi_helper", HELPER)
b = importlib.util.module_from_spec(spec)
assert spec and spec.loader
spec.loader.exec_module(b)
b.SHOT = SHOT


def configure(doc: Document):
    b.configure_document(doc)
    header = doc.sections[0].header.tables[0]
    b.set_cell_text(
        header.cell(0, 0),
        "eBisnis POS  •  KEUANGAN → AKUNTANSI",
        bold=True,
        color=b.BLUE,
        size=8,
    )
    b.set_cell_text(
        header.cell(0, 1),
        "Manual Pengguna & Hasil UAT",
        bold=True,
        color=b.MUTED,
        size=8,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
    )
    footer = doc.sections[0].footer.paragraphs[0]
    if footer.runs:
        footer.runs[0].text = "Dokumen UAT • Data demo • 4 September 2026"


def page_break(doc):
    doc.add_page_break()


def cover(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    tag = doc.add_table(rows=1, cols=1)
    b.shade(tag.cell(0, 0), b.SKY)
    b.set_cell_text(
        tag.cell(0, 0),
        "BUILD VARIAN eBISNIS  •  UAT 04-09-2026  •  VERSI 1.34.20 (BUILD 182)",
        bold=True,
        color=b.BLUE,
        size=9,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("MANUAL PENGGUNA & HASIL UAT")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(17)
    r.font.color.rgb = RGBColor.from_string(b.BLUE)

    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(5)
    p.add_run("Keuangan → Akuntansi")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        "Uang muka, kas, transfer, pajak, jurnal otomatis, posting, dan laporan keuangan"
    )
    r.font.name = "Aptos Display"
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor.from_string(b.MUTED)

    b.add_screenshot(
        doc,
        "01-menu-keuangan-terbuka.png",
        "Gambar 1. Grup Keuangan pada aplikasi Windows eBisnis—render asli, bukan mockup.",
        width=6.35,
    )
    rows = (
        ("Produk", "eBisnis POS — varian eBisnis"),
        ("Versi diuji", "1.34.20 (build 182)"),
        ("Lingkungan", "Server demo eBisnis • Semua Toko"),
        ("Periode laporan", "1–30 September 2026"),
        ("Keluaran", "Panduan pengguna + matriks UAT + 37 screenshot aktual"),
    )
    table = doc.add_table(rows=len(rows), cols=2)
    table.autofit = False
    table.columns[0].width = Cm(4.2)
    table.columns[1].width = Cm(13.0)
    for i, (key, value) in enumerate(rows):
        b.shade(table.cell(i, 0), b.NAVY)
        b.set_cell_text(table.cell(i, 0), key, bold=True, color=b.WHITE, size=8)
        b.shade(table.cell(i, 1), b.LIGHT)
        b.set_cell_text(table.cell(i, 1), value, size=8.3)
    page_break(doc)


def pair_page(doc, title, intro, first, second, callout=None):
    doc.add_heading(title, level=1)
    doc.add_paragraph(intro)
    b.add_screenshot(doc, first[0], first[1], width=6.10)
    b.add_screenshot(doc, second[0], second[1], width=6.10)
    if callout:
        b.add_callout(doc, *callout)
    page_break(doc)


def build():
    doc = Document()
    configure(doc)
    cover(doc)

    doc.add_heading("Cara menggunakan dokumen ini", level=1)
    doc.add_paragraph(
        "Dokumen ini menggabungkan panduan operasional dan hasil User Acceptance Test "
        "(UAT) seluruh submenu Keuangan, lalu menelusuri hasilnya ke Draft Jurnal dan "
        "laporan Akuntansi. Seluruh gambar berasal dari aplikasi Windows yang terhubung "
        "ke server demo eBisnis."
    )
    b.add_callout(
        doc,
        "Keamanan",
        "Kredensial demo tidak dicantumkan. Gunakan akses yang diberikan melalui kanal terpisah.",
        "info",
    )
    b.add_callout(
        doc,
        "Prinsip UAT",
        "Layar kosong akibat defect backend tetap ditampilkan dan diberi status BLOCKED/PARTIAL agar hasil dapat diaudit.",
        "warn",
    )
    doc.add_heading("Daftar isi ringkas", level=2)
    b.add_table(
        doc,
        ["Bagian", "Cakupan"],
        [
            ("1–2", "Alur end-to-end, data sampel, dan prasyarat"),
            ("3–10", "Uang Muka, Kas Besar, Kas Kecil, Dana Talangan, Reimbursement"),
            ("11–13", "Master Data, Proses Transfer, Penomoran Dokumen"),
            ("14", "Bayar Pajak"),
            ("15–16", "Draft Jurnal, posting otomatis, dan laporan"),
            ("17–18", "Matriks UAT, defect, retest, dan kontrol operasional"),
        ],
        widths=[3.0, 13.6],
    )
    page_break(doc)

    doc.add_heading("1. Alur proses end-to-end", level=1)
    doc.add_paragraph(
        "Transaksi dimulai dari pengajuan dan persetujuan, dilanjutkan pencairan atau "
        "transfer, pertanggungjawaban bila diperlukan, pembentukan jurnal, posting, lalu pelaporan."
    )
    b.add_table(
        doc,
        ["Tahap", "Menu", "Hasil yang diharapkan"],
        [
            ("1. Permintaan", "Uang Muka / Kas / Talangan / Reimbursement", "Dokumen bernomor dan siap disetujui"),
            ("2. Persetujuan", "Daftar transaksi", "Status disetujui; nilai dan penerima tervalidasi"),
            ("3. Pencairan", "Proses Transfer", "DPC direalisasikan dan bukti transfer tercatat"),
            ("4. Akuntabilitas", "Pertanggungjawaban / Penggantian", "Rincian biaya dan pengembalian terdokumentasi"),
            ("5. Jurnal", "Akuntansi → Draft Jurnal", "Debet/kredit otomatis tampil per jenis aktivitas"),
            ("6. Pelaporan", "Akuntansi → Katalog Laporan", "Transaksi terposting masuk laporan keuangan"),
        ],
        widths=[3.0, 5.0, 8.6],
    )
    doc.add_heading("Kontrol utama", level=2)
    b.add_steps(
        doc,
        [
            "Pastikan Satuan Kerja, tanggal, penerima, dan jenis transaksi benar.",
            "Gunakan judul yang menyebut tujuan kegiatan dan periode.",
            "Pisahkan pembuat, penyetuju, dan pelaksana transfer bila memungkinkan.",
            "Posting hanya setelah akun master dan bukti transaksi lengkap.",
        ],
    )
    b.add_callout(
        doc,
        "Hasil",
        "Rangkaian UAT menghasilkan transaksi nyata, 30 aktivitas terposting, dan laporan bernilai—bukan kumpulan layar kosong.",
        "ok",
    )
    page_break(doc)

    doc.add_heading("2. Inventaris data sampel UAT", level=1)
    b.add_table(
        doc,
        ["Objek", "Data akhir", "Contoh praktik baik"],
        [
            ("Uang Muka", "3 dokumen • Rp4.650.000", "Operasional September; Pelatihan Tim"),
            ("PJ Uang Muka", "2 dokumen • Rp2.650.000", "ATK, transportasi, konsumsi; kembali Rp200.000"),
            ("Kas Besar", "3 dokumen • Rp2.700.089", "Perawatan fasilitas; servis AC; material listrik"),
            ("PJ Kas Besar", "2 dokumen • Rp2.600.000", "Realisasi dan kembali Rp100.000"),
            ("Kas Kecil", "2 dokumen • Rp450.000", "Parkir, air minum, perlengkapan kebersihan"),
            ("Penggantian Kas Kecil", "2 dokumen • Rp450.000", "Pengisian kembali sesuai realisasi"),
            ("Dana Talangan", "1 dokumen • Rp750.000", "Talangan operasional yang dapat ditelusuri"),
            ("Reimbursement Pegawai", "0 dokumen baru", "Form diuji; penyimpanan BLOCKED oleh API"),
            ("Proses Transfer", "7 terealisasi • 1 menunggu", "Empat transfer sampel membentuk jurnal otomatis"),
            ("Pajak", "2 BAST PPN • Rp220.000", "Terutang tampil; penyetoran masih PARTIAL"),
        ],
        widths=[4.2, 4.6, 7.8],
        font_size=7.3,
    )
    doc.add_paragraph()
    b.add_callout(
        doc,
        "Master",
        "60 Jenis Pengeluaran, 9 Kategori Biaya Sales, 2 penerima Kas Besar, dan Reimbursement Tanpa Anggaran telah dipetakan ke akun.",
        "ok",
    )
    b.add_callout(
        doc,
        "Penomoran",
        "Sebelas alur dokumen memakai templat Standar Dokumen Keuangan eBisnis dengan contoh FIN/0000/IX/2026.",
        "ok",
    )
    page_break(doc)

    pages = [
        (
            "3. Uang Muka (Cash Advance)",
            "Uang Muka digunakan sebelum bukti biaya final tersedia. Dashboard menunjukkan tiga pengajuan senilai Rp4.650.000.",
            [
                "Pilih Keuangan → Uang Muka (Cash Advance).",
                "Periksa ringkasan jumlah, nilai, dan status.",
                "Klik Pengajuan Baru; pilih Satuan Kerja, jenis, dan penerima.",
                "Isi judul, nilai, tanggal kebutuhan, dan tujuan.",
                "Simpan, setujui, lalu cairkan melalui Proses Transfer.",
            ],
            "02-uang-muka-daftar.png",
            "Gambar 2. Dashboard Uang Muka: 3 pengajuan, total Rp4.650.000.",
            ("Praktik baik", "Judul harus spesifik; hindari keterangan umum tanpa periode dan tujuan.", "info"),
        ),
        (
            "3.1 Form Uang Muka",
            "Nomor dokumen mengikuti templat keuangan yang sudah dipasang.",
            [
                "Pastikan Satuan Kerja sesuai pemilik biaya.",
                "Pilih penerima yang akan mempertanggungjawabkan dana.",
                "Masukkan estimasi wajar dan dokumen pendukung.",
                "Gunakan status Diajukan untuk memulai persetujuan.",
            ],
            "03-uang-muka-formulir.png",
            "Gambar 3. Form Pengajuan Uang Muka Baru.",
            None,
        ),
        (
            "4. Pertanggungjawaban Uang Muka",
            "LPJ merinci pemakaian dana dan pengembalian sisa. Dua dokumen senilai Rp2.650.000 tersedia; pengembalian Rp200.000.",
            [
                "Pilih Keuangan → Pertanggungjawaban Uang Muka.",
                "Klik LPJ Baru dan pilih uang muka yang sudah dicairkan.",
                "Tambahkan baris ATK, transportasi, konsumsi, atau biaya sesuai bukti.",
                "Isi pengembalian bila realisasi lebih kecil.",
                "Simpan dan posting setelah seluruh akun terpetakan.",
            ],
            "04-pj-uang-muka-daftar.png",
            "Gambar 4. Dashboard LPJ Uang Muka dengan data nilai dan pengembalian.",
            ("Kontrol", "Total realisasi + pengembalian harus menjelaskan seluruh uang muka.", "warn"),
        ),
        (
            "4.1 Form LPJ Uang Muka",
            "Rincian biaya menjadi sumber pemilihan akun beban saat jurnal dibentuk.",
            [
                "Pilih dokumen yang belum memiliki LPJ.",
                "Gunakan Tambah Baris untuk setiap kategori biaya.",
                "Isi tanggal bukti, keterangan, nilai, dan jenis pengeluaran.",
                "Pastikan selisih nol atau dicatat sebagai pengembalian.",
            ],
            "05-pj-uang-muka-formulir.png",
            "Gambar 5. Form LPJ Uang Muka Baru.",
            None,
        ),
        (
            "5. Kas Besar",
            "Kas Besar menampung pengeluaran material atau berpersetujuan khusus. Tiga dokumen bernilai Rp2.700.089 tersedia.",
            [
                "Pilih Keuangan → Kas Besar.",
                "Klik Pengeluaran Baru dan pilih jenis serta Satuan Kerja.",
                "Isi penerima, nilai, tujuan, dan tanggal.",
                "Setelah disetujui, realisasikan melalui Proses Transfer.",
            ],
            "06-kas-besar-daftar.png",
            "Gambar 6. Dashboard Kas Besar: 3 dokumen dengan data bernilai.",
            ("Sampel", "Perawatan fasilitas dirinci menjadi servis AC dan material listrik.", "ok"),
        ),
        (
            "5.1 Form Kas Besar",
            "Form membedakan jenis kas, penerima, dan pemilik biaya.",
            [
                "Pilih jenis Kas Besar yang memiliki akun penerima.",
                "Isi judul spesifik, nilai, dan keterangan.",
                "Pilih Satuan Kerja yang menanggung biaya.",
                "Simpan sebagai Diajukan.",
            ],
            "07-kas-besar-formulir.png",
            "Gambar 7. Form Pengeluaran Kas Besar Baru.",
            None,
        ),
        (
            "6. Pertanggungjawaban Kas Besar",
            "Dua pertanggungjawaban senilai Rp2.600.000 tersedia. Sampel baru berhasil diposting dan mencatat kembali Rp100.000.",
            [
                "Pilih Keuangan → Pertanggungjawaban Kas Besar.",
                "Klik PJ Baru dan pilih Kas Besar yang direalisasikan.",
                "Masukkan baris realisasi berdasarkan bukti.",
                "Catat pengembalian dan periksa akun biaya.",
                "Posting setelah total dan referensi akun valid.",
            ],
            "08-pj-kas-besar-daftar.png",
            "Gambar 8. Dashboard Pertanggungjawaban Kas Besar.",
            ("Catatan", "Satu PJ lama gagal karena referensi akun warisan; PJ sampel baru berhasil.", "warn"),
        ),
        (
            "6.1 Form PJ Kas Besar",
            "Rincian biaya menjadi dasar debet akun beban dan penyelesaian saldo Kas Besar.",
            [
                "Pilih transaksi yang belum memiliki pertanggungjawaban.",
                "Tambahkan baris biaya dan jenis pengeluaran.",
                "Cocokkan total dengan kuitansi/invoice.",
                "Isi pengembalian bila ada sisa.",
            ],
            "09-pj-kas-besar-formulir.png",
            "Gambar 9. Form Pertanggungjawaban Kas Besar Baru.",
            None,
        ),
        (
            "7. Kas Kecil",
            "Kas Kecil dipakai untuk biaya rutin bernilai kecil. Dua dokumen Rp450.000 berisi parkir, air minum, dan kebersihan.",
            [
                "Pilih Keuangan → Kas Kecil.",
                "Klik Pengeluaran Baru.",
                "Pilih jenis pengeluaran dan isi nominal per bukti.",
                "Posting setelah akun biaya dan saldo kas kecil benar.",
            ],
            "10-kas-kecil-daftar.png",
            "Gambar 10. Dashboard Kas Kecil dengan dua transaksi sampel.",
            ("Praktik baik", "Gunakan Kas Kecil sesuai batas nominal kebijakan internal.", "info"),
        ),
        (
            "7.1 Form Kas Kecil",
            "Setiap baris perlu jenis pengeluaran agar jurnal memakai akun yang tepat.",
            [
                "Pilih Satuan Kerja dan jenis Kas Kecil.",
                "Tambahkan rincian biaya satu per satu.",
                "Isi tanggal dan uraian yang cocok dengan bukti.",
                "Periksa total sebelum menyimpan.",
            ],
            "11-kas-kecil-formulir.png",
            "Gambar 11. Form Pengeluaran Kas Kecil Baru.",
            None,
        ),
        (
            "8. Penggantian Kas Kecil",
            "Penggantian mengisi kembali saldo berdasarkan pengeluaran yang sudah diposting. Dua dokumen Rp450.000 tersedia.",
            [
                "Pilih Keuangan → Penggantian Kas Kecil.",
                "Klik Penggantian Baru dan pilih pengeluaran eligible.",
                "Cocokkan total dengan transaksi sumber.",
                "Transfer pengisian kembali lalu verifikasi jurnal.",
            ],
            "12-penggantian-kas-kecil-daftar.png",
            "Gambar 12. Dashboard Penggantian Kas Kecil.",
            ("Kontrol", "Jangan mengganti transaksi yang sudah pernah direimburse.", "warn"),
        ),
        (
            "8.1 Form Penggantian Kas Kecil",
            "Form memilih sumber Kas Kecil yang sah dan belum diganti.",
            [
                "Pilih Satuan Kerja dan rekening tujuan.",
                "Centang transaksi yang hendak diganti.",
                "Periksa total dan keterangan.",
                "Simpan, setujui, dan realisasikan transfer.",
            ],
            "13-penggantian-kas-kecil-formulir.png",
            "Gambar 13. Form Penggantian Kas Kecil Baru.",
            None,
        ),
        (
            "9. Dana Talangan",
            "Satu Dana Talangan Rp750.000 telah direalisasikan dan diposting.",
            [
                "Pilih Keuangan → Dana Talangan.",
                "Klik Talangan Baru dan tentukan Satuan Kerja serta penerima.",
                "Isi tujuan, nilai, tanggal, dan rencana penyelesaian.",
                "Setujui, transfer, dan pantau sampai diselesaikan.",
            ],
            "14-dana-talangan-daftar.png",
            "Gambar 14. Dashboard Dana Talangan dengan satu transaksi Rp750.000.",
            ("Kontrol", "Pantau umur talangan dan tindak lanjuti yang melewati jatuh tempo.", "warn"),
        ),
        (
            "9.1 Form Dana Talangan",
            "Uraian harus menghubungkan talangan dengan kegiatan, penerima, dan rencana pengembalian.",
            [
                "Pilih jenis talangan dan penerima.",
                "Masukkan nominal serta tanggal kebutuhan.",
                "Isi alasan bisnis dan rencana penyelesaian.",
                "Simpan untuk memulai persetujuan.",
            ],
            "15-dana-talangan-formulir.png",
            "Gambar 15. Form Dana Talangan Baru.",
            None,
        ),
        (
            "10. Reimbursement Pegawai — BLOCKED",
            "Layar dan formulir terbuka, tetapi penyimpanan contoh ditolak backend sehingga daftar tetap nol.",
            [
                "Pilih Keuangan → Reimbursement Pegawai.",
                "Klik Pengajuan Baru dan isi seluruh data bisnis.",
                "Tambahkan rincian biaya serta status Diajukan.",
                "Simpan hanya setelah perbaikan backend tersedia.",
            ],
            "16-reimbursement-pegawai-daftar.png",
            "Gambar 16. Dashboard kosong karena blocker penyimpanan.",
            ("Blocker API", "Insert gagal karena kolom wajib atasan tidak diisi oleh kontrak API saat ini.", "risk"),
        ),
        (
            "10.1 Form Reimbursement Pegawai",
            "UAT berhenti sebelum penyimpanan final agar tidak membuat data inkonsisten.",
            [
                "Pilih Satuan Kerja dan Jenis Reimbursement.",
                "Pilih Pegawai Penerima dan tanggal pengeluaran.",
                "Tambahkan rincian biaya dan keterangan.",
                "Sesudah backend diperbaiki, pastikan relasi atasan terisi valid.",
            ],
            "17-reimbursement-pegawai-formulir.png",
            "Gambar 17. Form Reimbursement Pegawai Baru.",
            None,
        ),
    ]
    for args in pages:
        b.add_page(
            doc,
            args[0],
            args[1],
            args[3],
            args[4],
            steps=args[2],
            callout=args[5],
        )

    doc.add_heading("11. Master Data Keuangan", level=1)
    doc.add_paragraph(
        "Master menentukan jenis transaksi, akun, metode transfer, dan kategori biaya. "
        "Kelengkapan master merupakan prasyarat posting otomatis."
    )
    b.add_steps(
        doc,
        [
            "Buka seluruh tab Master Data Keuangan.",
            "Pastikan setiap data aktif yang dipakai transaksi memiliki akun.",
            "Periksa jenis Uang Muka, Kas, Reimbursement, Pengeluaran, Transfer, dan Biaya Sales.",
            "Hindari mengganti akun setelah posting tanpa analisis dampak.",
        ],
    )
    b.add_screenshot(doc, "18-master-jenis-uang-muka.png", "Gambar 18. Master Data Keuangan—Jenis Uang Muka.", width=6.1)
    b.add_callout(doc, "Hasil", "Seed idempoten melengkapi akun tanpa membuat duplikasi.", "ok")
    page_break(doc)

    pair_page(
        doc,
        "11.1 Pemetaan akun biaya",
        "Jenis Pengeluaran dan Kategori Biaya Sales kini menunjukkan akun yang sesuai makna bisnis.",
        ("19-master-jenis-pengeluaran-terpetakan.png", "Gambar 19. Jenis Pengeluaran terpetakan ke akun."),
        ("20-master-kategori-biaya-sales-terpetakan.png", "Gambar 20. Kategori Biaya Sales terpetakan."),
        ("Hasil", "60 Jenis Pengeluaran dan 9 Kategori Biaya Sales tidak lagi memiliki akun wajib kosong.", "ok"),
    )

    b.add_page(
        doc,
        "12. Proses Transfer",
        "Dashboard menunjukkan 0 draft, 1 menunggu realisasi, 7 terealisasi, dan 1 DPC belum diproses.",
        "21-proses-transfer-dasbor.png",
        "Gambar 21. Dashboard Proses Transfer dengan status realisasi.",
        steps=[
            "Buka Proses Transfer dan pilih DPC yang sudah disetujui.",
            "Pilih rekening sumber/tujuan, tanggal, dan bukti.",
            "Simpan realisasi; sistem membentuk jurnal otomatis.",
            "Periksa tab Vendor dan Transitori bila digunakan.",
        ],
        callout=("Hasil UAT", "Empat transfer sampel baru masing-masing membentuk satu jurnal otomatis.", "ok"),
    )

    pair_page(
        doc,
        "12.1 Transfer dan pembayaran vendor",
        "Gunakan tab sesuai sumber DPC dan cocokkan nominal terhadap dokumen yang disetujui.",
        ("22-proses-transfer-terealisasi.png", "Gambar 22. Proses Transfer dan ringkasan 7 realisasi."),
        ("23-pembayaran-vendor.png", "Gambar 23. Tab Pembayaran Vendor."),
    )
    pair_page(
        doc,
        "12.2 Transitori",
        "Akun perantara harus direkonsiliasi sampai nilai masuk dan keluar selesai pada dokumen yang sama.",
        ("24-transitori-menunggu.png", "Gambar 24. Daftar Transitori Menunggu."),
        ("25-proses-transitori.png", "Gambar 25. Layar Proses Transitori."),
        ("Catatan", "Tab kosong karena seluruh DPC sampel direalisasikan langsung; ini bukan kegagalan pemuatan.", "info"),
    )

    pair_page(
        doc,
        "13. Penomoran Dokumen Keuangan",
        "Sebelas alur memakai templat Standar Dokumen Keuangan eBisnis sehingga tidak lagi jatuh ke barcode.",
        ("26-penomoran-alur-dokumen.png", "Gambar 26. Alur dokumen memakai templat standar."),
        ("27-penomoran-templat-standar.png", "Gambar 27. Templat FIN/0000/IX/2026."),
        ("Kontrol", "Uji satu dokumen baru dan pastikan nomor berikutnya konsisten.", "info"),
    )

    b.add_page(
        doc,
        "14. Bayar Pajak — PARTIAL",
        "Dua BAST PPN masing-masing Rp110.000 tampil. Penyetoran belum selesai karena defect backend.",
        "29-pajak-terutang.png",
        "Gambar 28. Pajak Terutang: 2 baris, total PPN Rp220.000.",
        steps=[
            "Buka Bayar Pajak → Pajak → Terutang.",
            "Pilih BAST dan cocokkan DPP, PPh, serta PPN.",
            "Sesudah perbaikan backend, klik Setor PPN/PPh.",
            "Verifikasi Riwayat Setoran dan jurnal pajak.",
        ],
        callout=("Status PARTIAL", "Penyetoran PPN gagal karena jenis_pajak_barang tidak diisi untuk baris PPN.", "risk"),
    )
    b.add_page(
        doc,
        "14.1 Riwayat Setoran Pajak",
        "Riwayat kosong merupakan konsekuensi langsung blocker setoran, bukan kekurangan eksplorasi UAT.",
        "30-pajak-riwayat-setoran.png",
        "Gambar 29. Riwayat Setoran kosong sebagai bukti defect.",
        steps=[
            "Perbaiki kontrak penyimpanan pajak.",
            "Setor satu BAST PPN.",
            "Pastikan nomor, tanggal, jenis, dan nilai tampil.",
            "Pastikan jurnal pajak terbentuk.",
        ],
        callout=("Retest", "Status dapat diubah menjadi LULUS hanya setelah riwayat dan jurnal pajak terverifikasi.", "warn"),
    )

    b.add_page(
        doc,
        "15. Integrasi ke Draft Jurnal",
        "Periode UAT menampilkan 81 aktivitas: 51 draft, 30 terposting, dan 0 closing.",
        "31-integrasi-draft-jurnal-semua.png",
        "Gambar 30. Ringkasan Draft Jurnal lintas modul.",
        steps=[
            "Pilih Akuntansi → Draft Jurnal.",
            "Atur periode yang mencakup transaksi Keuangan.",
            "Periksa kartu Draft, Terposting, Closing, dan Total.",
            "Buka kategori Uang Muka dan Kas serta Pengajuan Transfer.",
        ],
        callout=("Kontrol", "Tindak lanjuti draft yang gagal karena akun atau data warisan sebelum closing.", "warn"),
    )
    pair_page(
        doc,
        "15.1 Jurnal Uang Muka, Kas, dan Transfer",
        "Uang Muka dan Kas berisi 15 aktivitas—13 terposting. Pengajuan Transfer berisi 7 aktivitas dan seluruhnya terposting.",
        ("32-integrasi-jurnal-uang-muka-dan-kas.png", "Gambar 31. Uang Muka dan Kas: 13 terposting."),
        ("33-integrasi-jurnal-pengajuan-transfer.png", "Gambar 32. Pengajuan Transfer: 7 terposting."),
    )

    doc.add_heading("15.2 Pola jurnal yang diharapkan", level=1)
    doc.add_paragraph(
        "Nama akun aktual mengikuti konfigurasi Master Data. Tabel ini merupakan pola kontrol, "
        "bukan instruksi untuk memaksa kode akun tertentu."
    )
    b.add_table(
        doc,
        ["Peristiwa", "Debet", "Kredit", "Kontrol"],
        [
            ("Uang muka direalisasi", "Uang Muka", "Kas/Bank", "Nilai sesuai transfer"),
            ("LPJ diposting", "Beban terkait", "Uang Muka", "Rincian + kembali = uang muka"),
            ("Pengembalian sisa", "Kas/Bank", "Uang Muka", "Ada bukti setor kembali"),
            ("Kas Besar direalisasi", "Kas Besar/uang muka", "Kas/Bank", "Penerima dan rekening benar"),
            ("Kas Kecil diposting", "Beban terkait", "Kas Kecil", "Jenis pengeluaran terpetakan"),
            ("Penggantian Kas Kecil", "Kas Kecil", "Kas/Bank", "Tidak mengganti dua kali"),
            ("Dana Talangan", "Talangan/Piutang lain", "Kas/Bank", "Umur talangan dipantau"),
            ("Transfer DPC", "Akun tujuan", "Akun sumber", "Jurnal otomatis seimbang"),
            ("Setoran pajak", "Utang Pajak", "Kas/Bank", "Nomor dan jenis pajak"),
        ],
        widths=[4.0, 4.0, 4.0, 4.6],
        font_size=7.1,
    )
    b.add_callout(doc, "Kontrol posting", "Total debet harus sama dengan kredit dan referensi dokumen dapat ditelusuri.", "info")
    page_break(doc)

    reports = [
        (
            "16. Laporan Laba Rugi",
            "Laporan periode 1–30 September 2026 menampilkan Pendapatan Tidak Terikat Rp4.750.000.",
            "34-laporan-laba-rugi-integrasi-keuangan.png",
            "Gambar 33. Laba Rugi dengan nilai jurnal terposting.",
            [
                "Buka Akuntansi → Katalog Laporan.",
                "Cari Laba Rugi (Berbasis Jurnal Akuntansi).",
                "Isi periode 01–30 September 2026.",
                "Klik Tampilkan dan telusuri pendapatan, beban, serta laba.",
            ],
        ),
        (
            "16.1 Laporan Arus Kas",
            "Arus Kas menampilkan saldo awal Kas/Bank Rp12.549.911 dan mutasi menurut akun lawan.",
            "35-laporan-arus-kas-integrasi-keuangan.png",
            "Gambar 34. Arus Kas dengan saldo awal dan mutasi bernilai.",
            [
                "Cari Arus Kas (Berbasis Jurnal Akuntansi).",
                "Gunakan periode yang sama.",
                "Klik Tampilkan.",
                "Rekonsiliasi saldo awal, mutasi, dan saldo akhir.",
            ],
        ),
        (
            "16.2 Keseluruhan Jurnal",
            "Laporan memperlihatkan pasangan debet/kredit transfer Rp2.750.000, Rp750.000, dan Rp2.500.000.",
            "36-laporan-keseluruhan-jurnal.png",
            "Gambar 35. Jurnal transfer otomatis pada Keseluruhan Jurnal.",
            [
                "Cari Keseluruhan Jurnal (Jurnal Umum).",
                "Tampilkan periode September 2026.",
                "Telusuri nomor jurnal dan sumber transfer.",
                "Pastikan setiap jurnal seimbang.",
            ],
        ),
    ]
    for title, intro, shot, caption, steps in reports:
        callout = None
        if title.startswith("16.2"):
            callout = (
                "Catatan",
                "Endpoint daftar Jurnal Umum masih SERVER_ERROR, tetapi jalur laporan berhasil.",
                "warn",
            )
        b.add_page(doc, title, intro, shot, caption, steps=steps, callout=callout)

    doc.add_heading("17. Matriks hasil UAT seluruh submenu Keuangan", level=1)
    b.add_table(
        doc,
        ["Submenu", "Status", "Hasil"],
        [
            ("Uang Muka", "LULUS", "3 dokumen • Rp4.650.000; form dan dashboard tampil"),
            ("PJ Uang Muka", "LULUS", "2 dokumen • Rp2.650.000; kembali Rp200.000"),
            ("Kas Besar", "LULUS", "3 dokumen • Rp2.700.089"),
            ("PJ Kas Besar", "CATATAN", "Sampel baru lulus; 1 data lama gagal akun"),
            ("Kas Kecil", "LULUS", "2 dokumen • Rp450.000"),
            ("Penggantian Kas Kecil", "LULUS", "2 dokumen • Rp450.000"),
            ("Dana Talangan", "LULUS", "1 dokumen • Rp750.000"),
            ("Reimbursement Pegawai", "BLOCKED", "Form tampil; insert gagal field atasan"),
            ("Master Data Keuangan", "LULUS", "Pemetaan akun yang diuji lengkap"),
            ("Proses Transfer", "LULUS", "7 terealisasi; jurnal otomatis terverifikasi"),
            ("Penomoran Dokumen", "LULUS", "11 alur memakai templat FIN"),
            ("Bayar Pajak", "PARTIAL", "2 BAST Rp220.000; penyetoran gagal"),
        ],
        widths=[4.3, 2.8, 9.5],
        font_size=7.2,
    )
    doc.add_paragraph()
    b.add_callout(
        doc,
        "Kesimpulan",
        "9 submenu LULUS, 1 CATATAN data warisan, 1 BLOCKED, dan 1 PARTIAL. Integrasi jurnal serta tiga laporan utama berhasil.",
        "warn",
    )
    page_break(doc)

    doc.add_heading("18. Daftar defect dan retest", level=1)
    b.add_table(
        doc,
        ["ID", "Area", "Temuan", "Kriteria retest"],
        [
            ("FIN-01", "Reimbursement", "API tidak mengisi atasan wajib", "Tersimpan, disetujui, ditransfer, berjurnal"),
            ("FIN-02", "Bayar Pajak", "jenis_pajak_barang null", "Setoran, riwayat, dan jurnal tampil"),
            ("FIN-03", "PJ Kas Besar", "Data warisan gagal akun", "Data lama dapat diposting"),
            ("AKN-01", "Jurnal Umum", "jurnal_umum_daftar SERVER_ERROR", "Daftar membuka tanpa error"),
            ("UI-01", "Bayar Pajak", "Dua TabController memakai single ticker", "SUDAH DIPERBAIKI; tab membuka"),
        ],
        widths=[1.5, 3.2, 6.0, 5.9],
        font_size=7.2,
    )
    doc.add_heading("Urutan retest", level=2)
    b.add_steps(
        doc,
        [
            "Deploy perbaikan backend FIN-01 dan FIN-02.",
            "Buat Reimbursement Pegawai sampai transfer dan jurnal.",
            "Setor satu BAST PPN sampai Riwayat dan jurnal pajak.",
            "Koreksi data warisan PJ Kas Besar lalu posting.",
            "Retest daftar Jurnal Umum.",
            "Ulangi Laba Rugi dan Arus Kas.",
        ],
    )
    b.add_callout(
        doc,
        "Batas",
        "Closing dan Tutup Buku final tidak dijalankan karena dapat mengunci periode basis data demo bersama.",
        "warn",
    )
    page_break(doc)

    doc.add_heading("Lampiran. Bukti teknis dan checklist", level=1)
    b.add_table(
        doc,
        ["Pemeriksaan", "Hasil"],
        [
            ("Flutter Windows integration test", "Lulus; seluruh submenu dan laporan dibuka"),
            ("Finance test suites", "127 pengujian terkait lulus"),
            ("Seed idempoten", "Tidak menambah pemetaan/templat duplikat"),
            ("Screenshot", "37 PNG aplikasi aktual, 1264 × 681"),
            ("Varian build", "eBisnis POS 1.34.20 (build 182)"),
            ("Rahasia", "Kata sandi tidak disimpan dalam dokumen/gambar"),
        ],
        widths=[5.2, 11.4],
    )
    doc.add_heading("Checklist operasional", level=2)
    b.add_steps(
        doc,
        [
            "Akun master aktif dan terpetakan.",
            "Nomor dokumen berasal dari templat yang benar.",
            "Pembuat, penyetuju, dan transfer dapat ditelusuri.",
            "Bukti biaya/transfer dan tanggal konsisten.",
            "Debet sama dengan kredit sebelum posting.",
            "Draft sah diposting sebelum laporan/closing.",
            "Saldo laporan direkonsiliasi dengan jurnal.",
        ],
    )
    b.add_callout(
        doc,
        "Akhir dokumen",
        "Simpan bukti retest bersama manual ini agar status BLOCKED/PARTIAL dapat ditutup secara auditabel.",
        "info",
    )

    core = doc.core_properties
    core.title = "Manual UAT Keuangan dan Integrasi Akuntansi eBisnis"
    core.subject = "User manual dan hasil UAT build varian eBisnis"
    core.author = "Tim UAT eBisnis"
    core.comments = "Dibangun dari screenshot aktual aplikasi Windows eBisnis."
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
